from io import BytesIO

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models import Resume


class ExportService:
    def to_pdf(self, resume: Resume) -> bytes:
        stream = BytesIO()
        document = SimpleDocTemplate(stream, pagesize=LETTER, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
        styles = getSampleStyleSheet()
        content = resume.content or {}
        hidden = content.get("hidden_sections", [])
        story = [
            Paragraph(content.get("name") or resume.title, styles["Title"]),
            Paragraph(" · ".join([item for item in [content.get("email"), content.get("phone"), content.get("location")] if item]), styles["Normal"]),
            Spacer(1, 14),
        ]
        if "summary" not in hidden:
            story.extend([Paragraph("Summary", styles["Heading2"]), Paragraph(content.get("summary") or "Professional summary not added yet.", styles["BodyText"]), Spacer(1, 10)])
        if "skills" not in hidden:
            story.extend([Paragraph("Skills", styles["Heading2"]), Paragraph(", ".join(content.get("skills", [])) or "Skills not added yet.", styles["BodyText"])])
        for section in ["experience", "education", "links"]:
            if section in hidden:
                continue
            story.extend([Spacer(1, 10), Paragraph(section.title(), styles["Heading2"])])
            items = content.get(section, []) or ["No entries yet."]
            for item in items:
                story.append(Paragraph(f"• {item}", styles["BodyText"]))
        for custom in content.get("custom_sections", []):
            story.extend([Spacer(1, 10), Paragraph(custom.get("title", "Custom Section"), styles["Heading2"])])
            for item in custom.get("items", []):
                story.append(Paragraph(f"• {item}", styles["BodyText"]))
        document.build(story)
        return stream.getvalue()

    def to_docx(self, resume: Resume) -> bytes:
        document = Document()
        content = resume.content
        hidden = content.get("hidden_sections", [])
        document.add_heading(content.get("name") or resume.title, 0)
        if "summary" not in hidden:
            document.add_paragraph(content.get("summary", ""))
        if "skills" not in hidden:
            document.add_heading("Skills", level=1)
            document.add_paragraph(", ".join(content.get("skills", [])))
        for section in ["experience", "education", "links"]:
            if section in hidden:
                continue
            document.add_heading(section.title(), level=1)
            for item in content.get(section, []):
                document.add_paragraph(item, style="List Bullet")
        for custom in content.get("custom_sections", []):
            document.add_heading(custom.get("title", "Custom Section"), level=1)
            for item in custom.get("items", []):
                document.add_paragraph(item, style="List Bullet")
        stream = BytesIO()
        document.save(stream)
        return stream.getvalue()
